from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path.cwd()
ARTIFACT_DIR = ROOT / "artifacts" / "professor-guide"
SCREEN_DIR = ARTIFACT_DIR / "screenshots"
OUT_DOCX = ARTIFACT_DIR / "ai-cover-letter-professor-guide.docx"

FONT = "Malgun Gothic"
INK = RGBColor(10, 10, 18)
MUTED = RGBColor(95, 99, 109)
ACCENT = RGBColor(32, 37, 53)
BLUE = RGBColor(55, 92, 235)
LIGHT = "F7F7F8"
LINE = "D9D9DE"
DARK = "202532"


SLIDES = [
    {
        "page": "01",
        "slide": "00",
        "chapter": "Opening",
        "title": "AI 시대 자기소개서 작성법",
        "subtitle": "강의 목적과 오늘의 결과물을 한 화면에서 합의합니다.",
        "screenshot": "slide-01-00.png",
        "goal": "강의의 방향을 제시하고, 학생들이 오늘 만들 결과물이 '예쁜 글'이 아니라 제출 가능한 자기소개서 초안임을 이해하게 한다.",
        "script": [
            "오늘 강의는 자기소개서를 처음부터 문학적으로 쓰는 시간이 아닙니다. 이미 가진 경험을 회사가 이해할 수 있는 문장으로 재배치하는 시간입니다.",
            "AI는 대신 써주는 도구라기보다, 흩어진 경험을 항목별 질문 의도에 맞춰 정리해 주는 조력자로 다룹니다.",
            "수업 끝에는 학생 개인 정보, STAR 경험, 문장 교정 원칙을 반영해 Word 문서로 내려받을 수 있는 초안을 만드는 것이 목표입니다.",
            "첫 화면의 로봇은 전체 강의의 동반자 역할입니다. 마우스 커서에 반응하므로 도입부에서 가볍게 움직여 분위기를 열어도 좋습니다.",
        ],
        "how": [
            "첫 화면에서는 별도 입력 없이 강의 목표를 설명한다.",
            "오른쪽 화살표 또는 키보드 오른쪽 방향키로 다음 페이지로 이동한다.",
            "상단의 DOCX 버튼은 전체 흐름을 설명할 때만 언급하고, 실제 다운로드는 마지막 페이지에서 진행한다.",
        ],
        "emphasis": [
            "AI 사용의 핵심은 '대필'이 아니라 '구조화'라고 선명하게 말한다.",
            "학생들에게 오늘 필요한 자료는 개인 정보, 교육 과정, 프로젝트 경험, 배운 점이라고 안내한다.",
        ],
        "transition": "이제 강의를 이끄는 강사 프로필과 이 강의가 어떤 실무 관점에서 설계되었는지 확인하겠습니다.",
    },
    {
        "page": "02",
        "slide": "01",
        "chapter": "Instructor",
        "title": "강사 프로필",
        "subtitle": "강사의 경력과 수업 관점을 신뢰 요소로 제시합니다.",
        "screenshot": "slide-02-01.png",
        "goal": "강사의 현장 경험을 통해 강의 신뢰도를 만들고, 자기소개서가 채용 문서라는 관점을 먼저 세운다.",
        "script": [
            "이 페이지에서는 강사 이력 자체를 길게 자랑하기보다, 왜 이 수업이 '실무형 자기소개서'를 다루는지 연결해 주면 됩니다.",
            "핵심 문장은 화면 하단 인용문입니다. 영혼 없는 자기소개서의 문제는 스토리 부재가 아니라 배치와 해석의 부재입니다.",
            "학생이 같은 경험을 가지고 있어도 어떤 항목에 넣고, 어떤 의미로 해석하느냐에 따라 설득력이 달라집니다.",
            "따라서 오늘 수업은 경험을 인성, 태도, 직무역량으로 나누고 항목별 역할에 맞게 다시 조립하는 방식으로 진행됩니다.",
        ],
        "how": [
            "사진, 경력, 실습 도메인을 빠르게 훑고 인용문으로 시선을 모은다.",
            "로봇은 우측 상단에서 인터랙션을 유지하므로 강의 톤을 부드럽게 만드는 장치로 활용한다.",
            "인용문을 읽은 뒤 학생들에게 '내 경험은 어디에 넣어야 하는가'가 오늘의 질문이라고 안내한다.",
        ],
        "emphasis": [
            "스토리가 없어서 못 쓰는 것이 아니라, 스토리를 해석하지 못해서 약해지는 것이라고 강조한다.",
            "강사의 경력은 강의 권위를 위한 배경이고, 학생 활동으로 빠르게 넘어간다.",
        ],
        "transition": "그럼 왜 많은 자기소개서가 영혼 없어 보이는지, 그 원인을 먼저 보겠습니다.",
    },
    {
        "page": "03",
        "slide": "02",
        "chapter": "Why now",
        "title": "영혼 없는 자소서, 이제 그만",
        "subtitle": "막연한 경험 나열에서 벗어나 문장 판단 기준을 세웁니다.",
        "screenshot": "slide-03-02.png",
        "goal": "자기소개서 실패 원인을 학생 탓이 아니라 구조 부재로 재정의하고, 이후 실습의 필요성을 만든다.",
        "script": [
            "많은 학생이 자기소개서에서 '열심히 했다', '성실하다', '노력했다'를 반복합니다. 문제는 표현이 나쁜 것이 아니라 근거와 해석이 빠졌다는 점입니다.",
            "채용 담당자는 착한 문장을 찾는 것이 아니라, 이 사람이 우리 업무에서 어떤 행동을 할 사람인지 판단할 단서를 찾습니다.",
            "그래서 오늘은 경험을 시간순으로 늘어놓지 않고, 항목의 질문 의도에 맞춰 자르는 연습을 합니다.",
            "AI를 쓸 때도 같은 원칙이 적용됩니다. AI에게 많은 말을 주는 것보다, 어떤 역할의 문장을 만들지 지정하는 것이 중요합니다.",
        ],
        "how": [
            "문제 문장 예시를 구두로 하나 제시한다. 예: '저는 책임감이 강합니다.'",
            "학생에게 이 문장만으로는 왜 설득이 어려운지 질문한다.",
            "인성·태도와 직무역량을 분리해서 봐야 한다는 다음 페이지의 큰 틀로 연결한다.",
        ],
        "emphasis": [
            "나쁜 자기소개서는 경험이 없는 글이 아니라, 경험의 역할이 정리되지 않은 글이다.",
            "AI는 질문 의도와 기준을 넣었을 때 가장 잘 작동한다.",
        ],
        "transition": "이제 자기소개서의 대표 항목 5개를 지도로 놓고, 각 항목이 요구하는 역할을 보겠습니다.",
    },
    {
        "page": "04",
        "slide": "03",
        "chapter": "Framework",
        "title": "5개 카테고리 항목 지도",
        "subtitle": "성장과정, 장단점, 지원동기, 직무역량, 입사 후 포부의 역할을 나눕니다.",
        "screenshot": "slide-04-03.png",
        "goal": "자기소개서 항목을 한꺼번에 쓰려 하지 않고, 항목별로 다른 증거와 문장 역할을 갖는다는 점을 이해시킨다.",
        "script": [
            "자기소개서의 5개 항목은 같은 이야기를 반복하라고 있는 칸이 아닙니다. 각각 다른 질문을 하고 있습니다.",
            "성장과정은 현재 직무 태도의 뿌리를 보여주는 항목이고, 성격의 장단점은 태도와 보완 행동을 검증하는 항목입니다.",
            "직무 지원동기는 왜 이 직무를 선택했는지, 보유 직무역량은 실제 수행 가능성을, 입사 후 포부는 들어간 뒤 어떻게 기여할지를 묻습니다.",
            "학생들이 앞으로 입력하는 프로필과 STAR 경험은 이 5개 항목에 다시 배치될 재료입니다.",
        ],
        "how": [
            "카테고리 카드를 하나씩 짚으며 '목적', '근거', '문장 공식'을 설명한다.",
            "학생들에게 같은 경험이 여러 항목에 들어갈 수 있지만 해석은 달라져야 한다고 안내한다.",
            "하단 해설 영역은 선택된 항목의 쓰는 법과 피해야 할 방식을 보여주는 참고 영역으로 사용한다.",
        ],
        "emphasis": [
            "항목은 글자 수 칸이 아니라 채용 담당자의 질문이다.",
            "한 경험을 모든 항목에 그대로 복사하면 중복으로 보이므로, 역할을 바꿔 써야 한다.",
        ],
        "transition": "항목의 역할을 봤으니, 이제 실제 문장을 어떻게 고쳐야 하는지 문장 기준을 보겠습니다.",
    },
    {
        "page": "05",
        "slide": "04",
        "chapter": "Writing",
        "title": "나쁜 글의 11가지 / 좋은 문장 6원칙",
        "subtitle": "문장 첨삭 기준을 실시간 예시로 확인합니다.",
        "screenshot": "slide-05-04.png",
        "goal": "학생이 AI 첨삭을 받기 전에 어떤 표현이 감점되는지, 어떤 문장이 합격 문장에 가까운지 판단 기준을 갖게 한다.",
        "script": [
            "이 페이지는 문장 실험실입니다. 왼쪽에는 자주 감점되는 표현이 있고, 오른쪽에는 교정된 예시가 있습니다.",
            "자기소개서 문장은 멋있어 보이는 것보다 명확해야 합니다. 한 문장에 한 가지 생각만 담고, 주어와 행동을 분명히 세웁니다.",
            "학생 문장을 가운데 입력창에 붙여 넣고 AI 첨삭을 시작하면, 총평, 강점, 위험, 고친 예시가 순서대로 나옵니다.",
            "수업에서는 모든 학생이 AI 첨삭을 오래 기다리게 하기보다 대표 문장 하나를 넣고 기준을 함께 확인하는 방식이 좋습니다.",
        ],
        "how": [
            "왼쪽의 감점 표현 버튼을 눌러 예시 문장을 입력창에 불러온다.",
            "학생이 직접 고친 문장을 입력하게 한 뒤 AI 첨삭 시작 버튼을 누른다.",
            "오른쪽 Before/After 예시는 '어떻게 바뀌었는지'를 설명하는 비교 자료로 사용한다.",
        ],
        "emphasis": [
            "접속어 남용, 피동형, 번역투, 중복 표현은 문장의 힘을 약하게 만든다.",
            "AI 결과를 그대로 제출하지 말고, 왜 바뀌었는지 기준을 이해하게 한다.",
        ],
        "transition": "이제 문장 기준을 알았으니, AI가 초안을 만들 재료인 학생 프로필을 입력하겠습니다.",
    },
    {
        "page": "06",
        "slide": "05",
        "chapter": "[실습] Profile",
        "title": "학생 프로필 입력",
        "subtitle": "지원 직무, 교육 과정, 기술, 프로젝트 등 사실 정보를 채웁니다.",
        "screenshot": "slide-06-05.png",
        "goal": "자기소개서 생성에 필요한 사실 정보를 빠짐없이 채우고, AI가 허구를 만들지 않도록 입력 데이터의 기준을 세운다.",
        "script": [
            "여기부터는 실습입니다. 자기소개서 초안의 품질은 AI 모델보다 입력 정보의 정확도에 크게 좌우됩니다.",
            "성명, 연락처 같은 개인 신상은 문서 출력용 정보이고, 지원 직무, 지원 기업, 교육 과정, 기술 스택은 본문 생성에 영향을 줍니다.",
            "특히 훈련 내용과 프로젝트 설명은 학생이 실제로 배운 것과 해 본 일을 구체적으로 적어야 합니다.",
            "모르는 내용을 그럴듯하게 채우면 AI가 허구의 문장을 만들 수 있으니, 빈칸보다 거짓 정보가 더 위험하다고 설명합니다.",
        ],
        "how": [
            "개인 신상, 지원 정보, 교육/훈련, 보유 기술, 프로젝트 영역을 순서대로 입력한다.",
            "페이지가 길기 때문에 스크롤하며 모든 섹션을 확인한다.",
            "기술 스택은 쉼표로 구분하고, 실제 사용할 수 있는 기술만 남긴다.",
        ],
        "emphasis": [
            "프로필은 자기소개서의 재료 창고다. 여기에서 빠진 내용은 AI 초안에도 약하게 반영된다.",
            "지원 기업이 정해지지 않은 학생은 우선 희망 산업과 직무 중심으로 작성한다.",
        ],
        "transition": "프로필이 사실 정보라면, 다음은 자기소개서의 근거가 되는 경험을 STAR 구조로 입력하겠습니다.",
    },
    {
        "page": "07",
        "slide": "06",
        "chapter": "[실습] STAR",
        "title": "스토리(경험) 입력 - STAR",
        "subtitle": "상황, 과제, 행동, 결과, 배운 점과 정량 성과를 정리합니다.",
        "screenshot": "slide-07-06.png",
        "goal": "학생의 경험을 자기소개서 문장으로 바꾸기 전에, 상황과 행동과 결과를 분리해 증거 형태로 정리한다.",
        "script": [
            "STAR는 경험을 설득 가능한 증거로 바꾸는 도구입니다. 상황은 배경, 과제는 해결해야 했던 문제, 행동은 본인이 실제로 한 일, 결과는 변화입니다.",
            "많은 학생이 행동 칸에 팀 전체 성과를 쓰거나, 결과 칸에 느낀 점만 적습니다. 이 페이지에서는 칸마다 역할을 분리해야 합니다.",
            "정량 성과는 숫자가 있으면 좋지만, 없을 때는 기간 단축, 오류 감소, 반복 점검, 사용자 반응처럼 관찰 가능한 변화로 써도 됩니다.",
            "한 경험을 성장과정, 직무역량, 성격 장단점 등 여러 항목에 나눠 쓸 수 있으므로 제목과 분류도 함께 정리합니다.",
        ],
        "how": [
            "경험 추가 버튼으로 학생 경험을 2개 이상 입력하게 한다.",
            "각 경험마다 제목, 자기소개서 항목 분류, S/T/A/R, 배운 점, 정량 성과를 채운다.",
            "빈칸이 많으면 AI 초안이 추상적으로 나오므로 최소한 행동과 결과는 반드시 작성하게 한다.",
        ],
        "emphasis": [
            "자기소개서의 설득력은 경험의 크기가 아니라 행동의 구체성에서 나온다.",
            "팀 프로젝트는 '우리'가 아니라 '내가 맡은 역할'을 중심으로 쓴다.",
        ],
        "transition": "이제 프로필과 STAR 경험이 준비됐으니, AI가 5개 항목 초안을 한 번에 만들도록 하겠습니다.",
    },
    {
        "page": "08",
        "slide": "07",
        "chapter": "AI Studio",
        "title": "AI가 5개 항목을 한 번에 완성합니다",
        "subtitle": "프로필과 STAR 경험을 결합해 자기소개서 초안을 생성합니다.",
        "screenshot": "slide-08-07.png",
        "goal": "입력된 학생 정보를 바탕으로 자기소개서 5개 항목 초안을 생성하고, 생성 전 체크리스트를 확인한다.",
        "script": [
            "이 화면은 앞에서 입력한 프로필과 STAR 경험을 자기소개서 초안으로 변환하는 단계입니다.",
            "오른쪽 체크리스트는 AI가 초안을 만들기 전에 충분한 재료가 있는지 확인하는 장치입니다. 체크가 부족하면 문장이 일반론으로 흐를 가능성이 큽니다.",
            "자소서 초안 생성하기 버튼을 누르면 Review 페이지로 이동하며 항목별 초안이 생성됩니다.",
            "강의 중에는 버튼을 누르기 전에 학생들에게 'AI가 처음 만든 문장은 완성본이 아니라 1차 초안'이라는 점을 반드시 안내합니다.",
        ],
        "how": [
            "입력 요약과 생성 항목 5개를 확인한다.",
            "필수 입력이 충분하면 자소서 초안 생성하기 버튼을 누른다.",
            "생성 후 자동으로 Review 페이지로 넘어가므로, 결과가 나올 때까지 학생들에게 수정 기준을 다시 설명한다.",
        ],
        "emphasis": [
            "AI 초안은 제출물이 아니라 편집 가능한 시작점이다.",
            "체크리스트가 부족하면 이전 페이지로 돌아가 경험과 프로필을 보강한다.",
        ],
        "transition": "초안이 만들어지면 바로 제출하지 않고, 항목별로 첨삭하고 재작성하는 단계로 넘어갑니다.",
    },
    {
        "page": "09",
        "slide": "08",
        "chapter": "Review",
        "title": "섹션 첨삭 & 재작성",
        "subtitle": "AI 초안을 항목별로 읽고 직접 수정하며 다시 요청합니다.",
        "screenshot": "slide-09-08.png",
        "goal": "생성된 초안을 학생이 읽고, 항목별 목적에 맞게 제목과 본문을 직접 수정하도록 안내한다.",
        "script": [
            "여기서부터가 가장 중요합니다. AI가 만든 초안을 그대로 제출하는 것이 아니라, 학생 본인의 말과 실제 경험으로 다시 다듬어야 합니다.",
            "왼쪽 항목을 선택하면 성장과정, 성격의 장단점, 지원동기, 직무역량, 입사 후 포부를 차례로 검토할 수 있습니다.",
            "본문에서 어색한 표현, 과장된 말, 실제와 다른 내용은 반드시 직접 지웁니다. AI 재작성은 방향을 주는 도구이지 사실 확인을 대신하지 않습니다.",
            "교수자는 한 항목을 예로 들어 제목, 첫 문장, 근거 경험, 마지막 연결 문장을 함께 점검하면 좋습니다.",
        ],
        "how": [
            "항목 탭을 눌러 각 섹션의 제목과 본문을 확인한다.",
            "학생이 직접 문장을 고친 뒤 필요하면 항목별 AI 재작성을 호출한다.",
            "검수 포인트는 사실성, 항목 적합성, 중복 제거, 문장 자연스러움이다.",
        ],
        "emphasis": [
            "AI 문장 중 사실과 다른 내용은 무조건 삭제한다.",
            "각 항목의 첫 문장만 읽어도 답이 보여야 한다.",
        ],
        "transition": "검토가 끝났다면 마지막으로 Word 문서 형태로 내려받아 제출 가능한 파일로 정리하겠습니다.",
    },
    {
        "page": "10",
        "slide": "09",
        "chapter": "Export",
        "title": "DOCX 실무 산출",
        "subtitle": "작성 내용을 Word 문서로 내려받아 제출 전 최종 점검합니다.",
        "screenshot": "slide-10-09.png",
        "goal": "최종 자기소개서 초안을 Word 파일로 다운로드하고, 제출 전 수정 절차를 안내한다.",
        "script": [
            "마지막 페이지에서는 지금까지 입력하고 수정한 내용을 Word 문서로 내려받습니다.",
            "DOCX 버튼을 누르면 학생 이름 기준의 자기소개서 파일이 생성됩니다. 다운로드 후에는 반드시 Word에서 열어 양식, 줄바꿈, 글자 수, 맞춤법을 다시 확인해야 합니다.",
            "AI가 만들어 준 문장은 초안이므로 최종 제출 전에는 학생 본인의 표현으로 한 번 더 읽고 수정하는 과정이 필요합니다.",
            "교수자는 제출 전 체크리스트로 사실 확인, 중복 표현 제거, 회사명/직무명 확인, 문장 톤 통일을 안내하면 됩니다.",
        ],
        "how": [
            "상단 또는 화면 내 DOCX 버튼을 눌러 파일을 내려받는다.",
            "다운로드된 Word 파일을 열어 제목, 항목 순서, 본문 누락 여부를 확인한다.",
            "제출 전에는 학교 또는 회사 양식에 맞게 글자 크기와 여백을 최종 조정한다.",
        ],
        "emphasis": [
            "다운로드는 끝이 아니라 제출 전 최종 편집의 시작이다.",
            "AI 초안, 학생 검토, Word 최종 확인의 3단계를 지키게 한다.",
        ],
        "transition": "수업을 마무리하며 학생들이 각자 다운로드한 파일을 저장하고 다음 피드백 일정으로 연결합니다.",
    },
]


def set_font(run, size=None, bold=False, color=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_page_16_9(doc):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(13.333)
    section.page_height = Inches(7.5)
    section.top_margin = Inches(0.34)
    section.bottom_margin = Inches(0.34)
    section.left_margin = Inches(0.36)
    section.right_margin = Inches(0.36)
    section.header_distance = Inches(0.15)
    section.footer_distance = Inches(0.15)


def add_plain_para(cell_or_doc, text="", size=8.5, bold=False, color=INK, after=3, align=None):
    p = cell_or_doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, color=color)
    return p


def add_label(doc, text, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_font(run, size=7.8, bold=True, color=MUTED)
    return p


def add_section_title(cell, label, title):
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label)
    set_font(r1, size=7.3, bold=True, color=BLUE)
    r2 = p.add_run("  " + title)
    set_font(r2, size=9.8, bold=True, color=INK)


def add_bullets(cell, items, size=7.8):
    for item in items:
        p = cell.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1.4)
        p.paragraph_format.line_spacing = 1.03
        r = p.add_run(item)
        set_font(r, size=size, color=INK)


def add_cover(doc):
    add_label(doc, "AI COVER LETTER LECTURE · PROFESSOR GUIDE", WD_ALIGN_PARAGRAPH.CENTER)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run("AI 시대 자기소개서 작성법\n교수자 강의 스크립트 & 사용 가이드")
    set_font(run, size=24, bold=True, color=INK)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run("16:9 캡처 기반 · 슬라이드별 강의 멘트 · 실습 운영 방법 · Word 다운로드 안내")
    set_font(run, size=10.5, color=MUTED)

    table = doc.add_table(rows=1, cols=4)
    table.autofit = False
    set_table_borders(table, color="C9CBD2", size="6")
    headers = ["구성", "범위", "운영 방식", "최종 산출물"]
    values = [
        "10개 강의 화면",
        "Opening부터 Export까지",
        "슬라이드 캡처 + 교수자 스크립트",
        "학생 자기소개서 DOCX",
    ]
    for i, cell in enumerate(table.rows[0].cells):
        shade_cell(cell, LIGHT)
        set_cell_margins(cell, top=110, bottom=110, start=130, end=130)
        set_cell_border(cell)
        add_plain_para(cell, headers[i], size=7.4, bold=True, color=MUTED, after=2)
        add_plain_para(cell, values[i], size=10.2, bold=True, color=INK, after=0)

    doc.add_paragraph()
    intro = doc.add_table(rows=1, cols=2)
    intro.autofit = False
    set_table_borders(intro, color="D8DAE1", size="6")
    left, right = intro.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=120, bottom=120, start=150, end=150)
        set_cell_border(cell, color="D8DAE1", size="6")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    shade_cell(left, "FFFFFF")
    shade_cell(right, "F1F2F5")
    add_section_title(left, "USE", "문서 사용법")
    add_bullets(
        left,
        [
            "각 슬라이드는 캡처 페이지와 강의 노트 페이지로 구성되어 있습니다.",
            "캡처 페이지는 실제 강의 화면 확인용이고, 노트 페이지는 교수자 멘트와 조작 절차입니다.",
            "학생 실습 시에는 05 Profile, 06 STAR, 07 AI Studio, 08 Review, 09 Export 순서를 중점적으로 운영합니다.",
        ],
    )
    add_section_title(right, "FLOW", "강의 흐름")
    add_bullets(
        right,
        [
            "도입: AI 사용 목적과 강사 관점 형성",
            "이론: 항목 지도와 문장 교정 기준",
            "실습: 프로필 입력, STAR 경험 입력, AI 초안 생성",
            "마무리: 항목별 검토 후 DOCX 다운로드",
        ],
    )

    doc.add_page_break()


def add_capture_page(doc, slide):
    add_label(
        doc,
        f"PAGE {slide['page']} · SLIDE {slide['slide']} · {slide['chapter']}",
        WD_ALIGN_PARAGRAPH.LEFT,
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(slide["title"])
    set_font(r, size=18.5, bold=True, color=INK)

    img_path = SCREEN_DIR / slide["screenshot"]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.add_picture(str(img_path), width=Inches(10.8))

    doc.add_page_break()


def add_guide_page(doc, slide):
    header = doc.add_table(rows=1, cols=2)
    header.autofit = False
    left, right = header.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=60, bottom=70, start=100, end=100)
        set_cell_border(cell, color="FFFFFF", size="0")
    shade_cell(left, "FFFFFF")
    shade_cell(right, "FFFFFF")
    add_plain_para(left, f"{slide['page']} · {slide['chapter']}", size=7.6, bold=True, color=MUTED, after=1)
    add_plain_para(left, slide["title"], size=16.2, bold=True, color=INK, after=1)
    add_plain_para(left, slide["subtitle"], size=8.4, color=MUTED, after=0)
    add_plain_para(right, "강의 목표", size=7.6, bold=True, color=BLUE, after=1, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_plain_para(right, slide["goal"], size=8.1, color=INK, after=0, align=WD_ALIGN_PARAGRAPH.RIGHT)

    main = doc.add_table(rows=1, cols=2)
    main.autofit = False
    set_table_borders(main, color="D4D6DD", size="6")
    left, right = main.rows[0].cells
    for cell in (left, right):
        set_cell_margins(cell, top=120, bottom=100, start=140, end=140)
        set_cell_border(cell, color="D4D6DD", size="6")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    shade_cell(left, "FFFFFF")
    shade_cell(right, "F5F6F8")
    add_section_title(left, "SCRIPT", "교수자 멘트")
    for paragraph in slide["script"]:
        add_plain_para(left, paragraph, size=8.15, color=INK, after=3)

    add_section_title(right, "HOW", "사용 방법")
    add_bullets(right, slide["how"], size=7.7)
    add_section_title(right, "POINT", "강조 포인트")
    add_bullets(right, slide["emphasis"], size=7.7)
    add_section_title(right, "NEXT", "전환 멘트")
    add_plain_para(right, slide["transition"], size=7.85, color=INK, after=0)


def build():
    ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    set_page_16_9(doc)

    styles = doc.styles
    styles["Normal"].font.name = FONT
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    styles["Normal"].font.size = Pt(9)

    add_cover(doc)
    for idx, slide in enumerate(SLIDES):
        add_capture_page(doc, slide)
        add_guide_page(doc, slide)
        if idx != len(SLIDES) - 1:
            doc.add_page_break()

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build()
